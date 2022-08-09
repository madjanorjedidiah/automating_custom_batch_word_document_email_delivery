from docx import Document
import re
import csv
import smtplib
from os.path import basename
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import os


# edit the document in question and return a customised document with the users name
def edit_docx(file_path, file, names, addresses, contacts):
	curr_file = os.path.join(file_path, file)
	document = Document(curr_file)
	# change the address, name, contact to suit your use case thus the placeholders in the document you wish to edit
	address = re.compile("ADDRESS")
	name = re.compile("NAME")
	contact = re.compile("CONTACT")
	a = 0
	for p in document.paragraphs:
		a += 1
		new_doc = os.path.join(file_path, names + '.docx')
		if address.search(p.text):
			p.text = p.text.replace('ADDRESS', addresses)
			document.save(new_doc)
		if name.search(p.text):
			p.text = p.text.replace('NAME', names)
			document.save(new_doc)
		if contact.search(p.text):
			p.text = p.text.replace('CONTACT', '0'+ f'{contacts}')
			document.save(new_doc)
		if a == len(document.paragraphs):
			break
	print("done")
	return (names, new_doc)


#  send email with attached document generated
def send_mail(send_email_from, send_email_to, subject, text_to_sender, file_name, file=None):
	msg = MIMEMultipart()
	msg['From'] = send_email_from
	msg['To'] = send_email_to
	msg['Subject'] = subject

	msg.attach(MIMEText(text_to_sender))

	ff = open(file, "rb")
	part = MIMEApplication(
	    ff.read(),
	    Name=basename(file)
	)
	# After the file is closed
	part['Content-Disposition'] = 'attachment; filename="%s"' % file_name + '.docx'
	msg.attach(part)
	server = smtplib.SMTP('smtp.gmail.com', 587)
	server.starttls()
	server.login(send_email_from, put_string_password_here)
	server.sendmail(send_email_from, send_email_to, msg.as_string())
	server.close()
	return 'email_sent'


def doc_automation(csv_file_path, csv_file_name, doc_file_path, doc_file_name, send_email_from, send_email_to, subject, text_to_sender):
	# get the csv document and read the content basically the details of the people needed to populate the word template.
	csv_full_file = os.path.join(csv_file_path, csv_file_name)
	csv_doc = open(csv_full_file, "r")
	csv_reader = csv.reader(csv_doc)     # read the csv
	next(csv_reader)   # skip the header

	for row in csv_reader:
		file = edit_docx(doc_file_path, doc_file_name, row[0], row[2], row[1]) # change the row indexes to suit your use case
		send_mail(send_email_from, send_email_to, subject, text_to_sender, file[0], file[1])
	return 'done'


# call the doc_automation to edit run the word customization and then send your emails